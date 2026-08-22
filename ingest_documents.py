import os
import re
import shutil
from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma


# ============================================================
# CONFIGURATION
# ============================================================

DOCUMENTS_DIR = Path("./documents")
CHROMA_DIR = Path("./chroma_db")

CHUNK_SIZE = 400       # approximate tokens
CHUNK_OVERLAP = 50     # approximately 12.5%


# ============================================================
# TEXT CLEANING
# ============================================================

def clean_text(text):
    if not text:
        return ""

    text = text.replace("\x00", " ")

    # Remove excessive whitespace
    text = re.sub(r"\s+", " ", text)

    return text.strip()


# ============================================================
# PDF PARSER
# ============================================================

def parse_pdf(file_path):

    print(f"\nReading PDF: {file_path.name}")

    reader = PdfReader(str(file_path))

    pages = []

    for page_number, page in enumerate(reader.pages, start=1):

        text = page.extract_text() or ""
        text = clean_text(text)

        if text:
            pages.append({
                "page": page_number,
                "text": text
            })

    print(f"Pages extracted: {len(pages)}")

    return pages


# ============================================================
# DOCX PARSER
# ============================================================

def parse_docx(file_path):

    print(f"\nReading DOCX: {file_path.name}")

    doc = DocxDocument(str(file_path))

    text_parts = []

    for paragraph in doc.paragraphs:

        text = clean_text(paragraph.text)

        if text:
            text_parts.append(text)

    full_text = " ".join(text_parts)

    if not full_text:
        return []

    return [{
        "page": 1,
        "text": full_text
    }]


# ============================================================
# CHUNKING
# ============================================================

def create_chunks(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):

    words = text.split()

    chunks = []

    start = 0

    while start < len(words):

        end = start + chunk_size

        chunk_words = words[start:end]

        if chunk_words:
            chunks.append(" ".join(chunk_words))

        if end >= len(words):
            break

        start = end - overlap

    return chunks


# ============================================================
# DOCUMENT PROCESSING
# ============================================================

def process_file(file_path):

    extension = file_path.suffix.lower()

    if extension == ".pdf":
        pages = parse_pdf(file_path)

    elif extension == ".docx":
        pages = parse_docx(file_path)

    else:
        return []

    documents = []

    global_chunk_number = 0

    for page_data in pages:

        page_number = page_data["page"]
        page_text = page_data["text"]

        chunks = create_chunks(page_text)

        for chunk_number, chunk in enumerate(chunks):

            chunk_id = (
                f"{file_path.stem}_"
                f"page_{page_number}_"
                f"chunk_{chunk_number}"
            )

            document = Document(
                page_content=chunk,
                metadata={
                    "document": file_path.name,
                    "source": file_path.name,
                    "page": page_number,
                    "section": "Unknown",
                    "chunk_id": chunk_id,
                }
            )

            documents.append(document)

            global_chunk_number += 1

    print(
        f"Created {global_chunk_number} chunks "
        f"from {file_path.name}"
    )

    return documents


# ============================================================
# LOAD ALL DOCUMENTS
# ============================================================

def load_documents():

    all_documents = []

    supported_extensions = {".pdf", ".docx"}

    files = [
        file
        for file in DOCUMENTS_DIR.iterdir()
        if file.is_file()
        and file.suffix.lower() in supported_extensions
    ]

    if not files:
        print("No PDF or DOCX documents found.")
        return []

    print("\n==============================================")
    print("DOCUMENT INGESTION")
    print("==============================================")

    for file_path in files:

        documents = process_file(file_path)

        all_documents.extend(documents)

    return all_documents


# ============================================================
# BUILD CHROMADB
# ============================================================

def build_vector_store(documents):

    print("\n==============================================")
    print("BUILDING CHROMADB")
    print("==============================================")

    print(f"Total chunks: {len(documents)}")

    print("\nLoading embedding model...")

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

    # Remove old database so duplicate vectors are not retained.
    if CHROMA_DIR.exists():

        print("\nRemoving old ChromaDB...")

        shutil.rmtree(CHROMA_DIR)

    print("\nCreating fresh ChromaDB...")

    vector_db = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory=str(CHROMA_DIR),
    )

    print("\n==============================================")
    print("CHROMADB CREATED SUCCESSFULLY")
    print("==============================================")

    print(f"Database: {CHROMA_DIR}")
    print(f"Stored chunks: {len(documents)}")

    return vector_db


# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":

    print("\n==============================================")
    print("SRA DOCUMENT INGESTION PIPELINE")
    print("==============================================")

    documents = load_documents()

    if not documents:

        print("\nNo documents were processed.")
        raise SystemExit(1)

    print(
        f"\nTotal documents/chunks prepared: "
        f"{len(documents)}"
    )

    build_vector_store(documents)

    print("\n==============================================")
    print("INGESTION COMPLETED")
    print("==============================================")